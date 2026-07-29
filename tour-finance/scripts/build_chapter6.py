# -*- coding: utf-8 -*-
"""สร้างบท 6 (ระเบียบวิธีวิจัย) ฉบับแก้สมบูรณ์ เป็นไฟล์ .docx พร้อมวางลงเล่ม
รันด้วย: python scripts/build_chapter6.py
ผลลัพธ์: output/บทที่6_ระเบียบวิธีวิจัย_ฉบับแก้.docx
ฟอนต์: TH Sarabun New (มาตรฐานวิชาการไทย) ตัวอักษร 16
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "TH Sarabun New"
BODY = 16      # ขนาดตัวอักษรเนื้อหา
H1 = 18        # หัวข้อใหญ่ 6.x
H2 = 16        # หัวข้อย่อย 6.x.y

doc = Document()

# ---- ตั้งค่าฟอนต์เริ่มต้นของสไตล์ Normal (รวม complex-script สำหรับภาษาไทย) ----
def set_style_font(style, size, bold=False):
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), FONT)
    # ขนาดสำหรับ complex script (ไทย)
    szcs = rpr.find(qn('w:szCs'))
    if szcs is None:
        szcs = OxmlElement('w:szCs'); rpr.append(szcs)
    szcs.set(qn('w:val'), str(size * 2))

normal = doc.styles['Normal']
set_style_font(normal, BODY)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

def _apply_run_font(run, size, bold, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.append(rfonts)
    for a in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(a), FONT)
    szcs = rpr.find(qn('w:szCs'))
    if szcs is None:
        szcs = OxmlElement('w:szCs'); rpr.append(szcs)
    szcs.set(qn('w:val'), str(size * 2))
    if bold:
        bcs = rpr.find(qn('w:bCs'))
        if bcs is None:
            bcs = OxmlElement('w:bCs'); rpr.append(bcs)

def heading(text, size, level=None):
    p = doc.add_paragraph()
    if level is not None:
        p.style = doc.styles[f'Heading {level}']
    r = p.add_run(text)
    _apply_run_font(r, size, True, RGBColor(0x1a, 0x1a, 0x1a))
    p.paragraph_format.space_before = Pt(10 if size >= H1 else 6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p

def para(text="", indent=True, bold=False, italic=False, size=BODY, align=None, space_after=6):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        _apply_run_font(r, size, bold)
        r.italic = italic
    return p

def bullet(text, size=BODY):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    _apply_run_font(r, size, False)
    p.paragraph_format.space_after = Pt(3)
    return p

def runs_para(segments, indent=True, space_after=6):
    """segments: list of (text, bold, italic)"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(space_after)
    for seg in segments:
        text, bold, italic = (seg + (False, False))[:3]
        r = p.add_run(text)
        _apply_run_font(r, BODY, bold)
        r.italic = italic
    return p

def formula(text):
    """สมการ — จัดกึ่งกลาง ตัวเอียง"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    _apply_run_font(r, BODY, False)
    r.italic = True
    return p

def shade_cell(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcpr.append(shd)

def make_table(headers, rows, widths, header_fill='D9E2F3'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); _apply_run_font(r, BODY, True)
        shade_cell(hdr[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(val); _apply_run_font(r, BODY, False)
            p.paragraph_format.space_after = Pt(2)
    # ตั้งความกว้างคอลัมน์
    for i, w in enumerate(widths):
        for row in t.rows:
            row.cells[i].width = Inches(w)
    return t

# =====================================================================
# เนื้อหา บท 6
# =====================================================================
heading("6. ระเบียบวิธีวิจัย", H1, 1)
para("การวิจัยครั้งนี้เป็นการวิจัยเชิงปริมาณ (Quantitative Research) โดยใช้ข้อมูลทุติยภูมิแบบแผง (Panel Data) "
     "รายไตรมาสของบริษัทจดทะเบียนกลุ่มธุรกิจโรงแรมในตลาดหลักทรัพย์แห่งประเทศไทย เพื่อเปรียบเทียบประสิทธิภาพ"
     "การบริหารเงินทุนหมุนเวียน (วงจรเงินสด) และความสามารถในการทำกำไร (ROA, ROE) ระหว่างกลุ่มที่กระจายความ"
     "เสี่ยงไปต่างประเทศกับกลุ่มที่ดำเนินธุรกิจเฉพาะภายในประเทศ โดยมีรายละเอียดของระเบียบวิธีดังนี้")

# ---- 6.1 ----
heading("6.1 ประชากรและกลุ่มตัวอย่าง", H1, 2)
para("ประชากรที่ใช้ในการศึกษา คือ บริษัทจดทะเบียนในตลาดหลักทรัพย์แห่งประเทศไทย (SET) กลุ่มอุตสาหกรรมบริการ "
     "หมวดธุรกิจการท่องเที่ยวและสันทนาการ (Tourism & Leisure) ที่ประกอบธุรกิจโรงแรมเป็นหลัก ผู้วิจัยเลือกกลุ่ม"
     "ตัวอย่างแบบเจาะจง (Purposive Sampling) ตามเกณฑ์การคัดกรอง ดังนี้")
bullet("มีสัดส่วนรายได้หลักจากการประกอบธุรกิจโรงแรมหรือรับบริหารจัดการโรงแรมไม่น้อยกว่าร้อยละ 50 ของรายได้รวม")
bullet("มีสถานะการดำเนินงานและการซื้อขายปกติ ไม่ถูกขึ้นเครื่องหมายห้ามซื้อขาย (SP) หรือเข้าข่ายอาจถูกเพิกถอน (NC) "
       "ตลอดช่วงเวลาที่ศึกษา และมีส่วนของผู้ถือหุ้นเป็นบวกในทุกงวดเพื่อความถูกต้องในการคำนวณอัตราส่วน")
bullet("มีรอบระยะเวลาบัญชีสิ้นสุด ณ วันที่ 31 ธันวาคม ของทุกปี เพื่อควบคุมผลกระทบจากฤดูกาล (Seasonality) "
       "และให้เปรียบเทียบข้ามบริษัทได้บนงวดเวลาเดียวกัน")
bullet("มีข้อมูลงบการเงินและรายงานประจำปีครบถ้วนตลอด 5 ปี (พ.ศ. 2564–2568) ครบทั้ง 4 ไตรมาสในแต่ละปี")
para("จากเกณฑ์ดังกล่าว มีบริษัทผ่านการคัดเลือกจำนวน 10 บริษัท ผู้วิจัยจำแนกกลุ่มตามกลยุทธ์การกระจายความเสี่ยง"
     "ทางภูมิศาสตร์ โดยใช้เกณฑ์ “สัดส่วนรายได้จากการดำเนินงานในต่างประเทศตั้งแต่ร้อยละ 20 ของรายได้รวมขึ้นไป” "
     "ซึ่งพิจารณาจากข้อมูลรายได้จำแนกตามส่วนงานภูมิศาสตร์ (Geographical Segment) ในหมายเหตุประกอบงบการเงิน"
     "และรายงานประจำปี (แบบ 56-1 One Report) ของแต่ละบริษัท แบ่งได้ 2 กลุ่ม ดังตารางที่ 6.1")
para("ตารางที่ 6.1 การจำแนกกลุ่มตัวอย่างตามกลยุทธ์การกระจายความเสี่ยงทางภูมิศาสตร์", indent=False, bold=True, space_after=3)
make_table(
    ["กลุ่ม", "จำนวน", "บริษัท (ตัวย่อ)"],
    [
        ["กระจายความเสี่ยงต่างประเทศ (DIV = 1)", "3", "MINT, SHR, DUSIT"],
        ["ดำเนินธุรกิจในประเทศ (DIV = 0, กลุ่มฐาน)", "7", "ERW, VRANDA, OHTL, SHANG, ASIA, MANRIN, CENTEL"],
    ],
    widths=[2.6, 0.7, 3.4],
)
runs_para([("หมายเหตุ: ", True, False),
           ("การจำแนกกลุ่มยึดสัดส่วนรายได้ต่างประเทศจากหมายเหตุส่วนงานเป็นหลัก ผู้วิจัยจะรายงานสัดส่วน"
            "รายได้ต่างประเทศของแต่ละบริษัทประกอบไว้ในภาคผนวก เพื่อยืนยันการจัดกลุ่มอย่างโปร่งใส ทั้งนี้ กรณีที่บริษัท"
            "มีการดำเนินงานในต่างประเทศแต่สัดส่วนรายได้ต่ำกว่าร้อยละ 20 (เช่น CENTEL) จะจัดอยู่ในกลุ่มดำเนินธุรกิจในประเทศ", False, False)],
          space_after=8)

# ---- 6.2 ----
heading("6.2 การเก็บรวบรวมข้อมูล", H1, 2)
para("งานวิจัยนี้ใช้ข้อมูลทุติยภูมิ (Secondary Data) เชิงปริมาณในรูปแบบข้อมูลแบบแผง (Panel Data) โดยใช้ข้อมูล"
     "รายไตรมาส (Quarterly Data) ตั้งแต่ไตรมาสที่ 1 ปี พ.ศ. 2564 ถึงไตรมาสที่ 4 ปี พ.ศ. 2568 รวม 5 ปี (20 ไตรมาส) "
     "ของบริษัทกลุ่มตัวอย่าง 10 บริษัท รวมทั้งสิ้น 200 ค่าสังเกต (Observations) ซึ่งจัดเป็นข้อมูลแบบแผงที่มีมิติ"
     "ภาคตัดขวาง (Cross-section) 10 บริษัท และมิติเวลา (Time-series) 20 ไตรมาส")
para("ผู้วิจัยดึงตัวเลขทางการเงินจากงบแสดงฐานะการเงิน งบกำไรขาดทุนเบ็ดเสร็จ งบกระแสเงินสด หมายเหตุประกอบ"
     "งบการเงิน และรายงานประจำปี (แบบ 56-1 One Report) ที่เผยแพร่ผ่านตลาดหลักทรัพย์แห่งประเทศไทย (SET) และ"
     "ฐานข้อมูล SETSMART โดยข้อมูลที่ใช้คำนวณในแต่ละตัวแปรสรุปไว้ในหัวข้อ 6.3")
runs_para([("ข้อพึงระวังด้านความสม่ำเสมอของข้อมูล: ", True, False),
           ("(1) งบการเงินไตรมาสของไทยเป็นตัวเลขราย 3 เดือน ยกเว้นบางบริษัทขนาดเล็ก (เช่น MANRIN) ที่ยื่นงบแบบ"
            "สะสม (Year-to-Date) และไม่มีงบรายปีแยกต่างหาก ผู้วิจัยจึงแปลงเป็นตัวเลขรายไตรมาสด้วยวิธีหักลบยอดสะสม"
            "งวดก่อนหน้า (De-cumulation) และสังเคราะห์งบรายปีจากผลรวม 4 ไตรมาส (2) ตัวเลขไตรมาสที่ 4 คำนวณจาก "
            "งบรายปีหักด้วยไตรมาสที่ 1–3 (3) มีการตรวจสอบดุลบัญชี (สินทรัพย์รวม = หนี้สินรวม + ส่วนของผู้ถือหุ้น) "
            "และกระทบยอดกำไรรวม = ส่วนผู้ถือหุ้นใหญ่ + ส่วนได้เสียที่ไม่มีอำนาจควบคุม ในทุกงวด", False, False)],
          space_after=8)

# ---- 6.3 ----
heading("6.3 ตัวแปรที่ใช้ในการวิจัยและการวัดค่า", H1, 2)
para("งานวิจัยนี้ประกอบด้วยตัวแปรตาม 3 ตัว ตัวแปรอิสระ 1 ตัว และตัวแปรควบคุม (รวมตัวแปรควบคุมด้านฤดูกาลและ"
     "ช่วงวิกฤต) โดยมีนิยามเชิงปฏิบัติการและสูตรการวัดค่าดังนี้")

heading("6.3.1 ตัวแปรอิสระ (Independent Variable)", H2, 3)
runs_para([("การกระจายความเสี่ยงไปต่างประเทศ (DIV): ", True, False),
           ("เนื่องจากกลยุทธ์การกระจายความเสี่ยงเป็นข้อมูลเชิงคุณภาพ ผู้วิจัยแปลงเป็นตัวแปรหุ่น (Dummy Variable) "
            "กำหนดค่าเท่ากับ 1 สำหรับกลุ่มที่มีสัดส่วนรายได้ต่างประเทศตั้งแต่ร้อยละ 20 ขึ้นไป และเท่ากับ 0 สำหรับกลุ่มที่"
            "ดำเนินธุรกิจเฉพาะในประเทศ (กลุ่มฐาน)", False, False)])
runs_para([("ข้อสังเกตเชิงระเบียบวิธี: ", True, True),
           ("ตัวแปร DIV เป็นตัวแปรที่ไม่เปลี่ยนแปลงตามเวลา (Time-invariant) เนื่องจากสถานะกลุ่มของแต่ละบริษัทคงที่"
            "ตลอด 20 ไตรมาส ซึ่งมีนัยต่อการเลือกแบบจำลอง (ดูหัวข้อ 6.4.4) นอกจากนี้ ผู้วิจัยจะเก็บสัดส่วนรายได้ต่างประเทศ"
            "แบบต่อเนื่อง (ร้อยละ) ไว้เป็นตัวแปรทางเลือกเพื่อการทดสอบความคงทน (Robustness) ของผลการศึกษาด้วย", False, False)],
          space_after=8)

heading("6.3.2 ตัวแปรตาม (Dependent Variables)", H2, 3)
para("เนื่องจากข้อมูลเป็นรายไตรมาส ผู้วิจัยคำนวณอัตราส่วนบนฐานรายไตรมาสอย่างสอดคล้องกับงวดของกระแส (Flow) "
     "โดยใช้ยอดงบดุลแบบค่าเฉลี่ยต้นงวด–ปลายงวด และจำนวนวันตามปฏิทินจริงของแต่ละไตรมาส (ไตรมาส 1 = 90 วัน, "
     "ไตรมาส 2 = 91 วัน, ไตรมาส 3 = 92 วัน, ไตรมาส 4 = 92 วัน) ดังนี้")

runs_para([("1) วงจรเงินสด (Cash Conversion Cycle: CCC) ", True, False),
           ("วัดประสิทธิภาพการบริหารเงินทุนหมุนเวียน คำนวณจาก", False, False)])
formula("CCC = DIO + DSO − DPO")
para("โดยตัวชี้วัดย่อยคำนวณบนฐานรายไตรมาส และใช้ยอด “ลูกหนี้การค้าและลูกหนี้อื่น” และ “เจ้าหนี้การค้าและเจ้าหนี้อื่น” "
     "ตามที่ปรากฏหน้างบแสดงฐานะการเงิน (Face of Balance Sheet) อย่างสม่ำเสมอทุกบริษัท เพื่อให้เปรียบเทียบข้าม"
     "บริษัทได้บนนิยามเดียวกัน (สอดคล้องกรอบ CFA Institute) ดังนี้", indent=True)
formula("DSO = [ ค่าเฉลี่ยลูกหนี้การค้าและอื่น ÷ รายได้ในไตรมาส ] × จำนวนวันในไตรมาส")
formula("DIO = [ ค่าเฉลี่ยสินค้าคงเหลือ ÷ ต้นทุนขายในไตรมาส ] × จำนวนวันในไตรมาส")
formula("DPO = [ ค่าเฉลี่ยเจ้าหนี้การค้าและอื่น ÷ ต้นทุนขายในไตรมาส ] × จำนวนวันในไตรมาส")
runs_para([("โดย ค่าเฉลี่ยยอดงบดุล = (ยอดต้นงวด + ยอดปลายงวด) ÷ 2 ", False, True),
           ("ทั้งนี้ แนวคิด CCC อ้างอิงต้นตำรับของ Richards & Laughlin (1980)", False, False)])

runs_para([("2) อัตราผลตอบแทนต่อสินทรัพย์รวม (Return on Assets: ROA) ", True, False),
           ("วัดความสามารถในการทำกำไรจากการใช้สินทรัพย์ คำนวณเป็นค่ารายไตรมาสจาก", False, False)])
formula("ROA = [ กำไรสุทธิรวมของงวด ÷ ค่าเฉลี่ยสินทรัพย์รวม ] × 100")
runs_para([("3) อัตราผลตอบแทนต่อส่วนของผู้ถือหุ้น (Return on Equity: ROE) ", True, False),
           ("วัดผลตอบแทนต่อผู้ถือหุ้น คำนวณเป็นค่ารายไตรมาสจาก", False, False)])
formula("ROE = [ กำไรสุทธิรวมของงวด ÷ ค่าเฉลี่ยส่วนของผู้ถือหุ้นรวม ] × 100")
runs_para([("หมายเหตุ: ", True, False),
           ("“กำไรสุทธิ” ใช้กำไรสุทธิรวม (Total Net Profit) และตัวหารใช้ยอดเฉลี่ยตามหลัก DuPont / Financial "
            "Statement Analysis โดย ROA และ ROE ในสมการหลักเป็นค่ารายไตรมาส (ไม่ปรับเป็นรายปี) เพื่อให้สอดคล้อง"
            "กับความถี่ของข้อมูลแบบแผง", False, False)], space_after=8)

heading("6.3.3 ตัวแปรควบคุม (Control Variables)", H2, 3)
runs_para([("1) ขนาดของกิจการ (Firm Size: SIZE) ", True, False),
           ("วัดจากลอการิทึมธรรมชาติของสินทรัพย์รวม: SIZE = ln(สินทรัพย์รวม)", False, False)])
runs_para([("2) โครงสร้างทางการเงิน (Leverage: LEV) ", True, False),
           ("วัดจากอัตราส่วนหนี้สินรวมต่อส่วนของผู้ถือหุ้น: LEV = หนี้สินรวม ÷ ส่วนของผู้ถือหุ้น", False, False)])
runs_para([("3) ตัวแปรหุ่นฤดูกาล (Seasonal Dummies: Q2, Q3, Q4) ", True, False),
           ("เนื่องจากธุรกิจโรงแรมมีความผันผวนตามฤดูกาลสูงและข้อมูลเป็นรายไตรมาส ผู้วิจัยกำหนดตัวแปรหุ่นไตรมาส "
            "Q2, Q3, Q4 (โดยให้ไตรมาสที่ 1 เป็นฐาน) เพื่อควบคุมผลกระทบเชิงฤดูกาล และลดปัญหาสหสัมพันธ์ในตัว", False, False)])
runs_para([("4) ตัวแปรหุ่นช่วงวิกฤตโรคระบาด (COVID) ", True, False),
           ("กำหนดค่าเท่ากับ 1 สำหรับไตรมาสที่ได้รับผลกระทบรุนแรงจากมาตรการควบคุมโรค (พ.ศ. 2564 ถึงต้นปี พ.ศ. 2565) "
            "และเท่ากับ 0 สำหรับงวดอื่น เพื่อควบคุมช่วงที่รายได้หดตัวผิดปกติซึ่งทำให้อัตราส่วนที่หารด้วยรายได้ (เช่น "
            "NPM, DSO, CCC) มีค่าผิดเพี้ยน และเพื่อรักษาสมมติฐานความแปรปรวนคงที่ของแบบจำลอง", False, False)], space_after=8)

# ---- 6.4 ----
heading("6.4 การวิเคราะห์ข้อมูล", H1, 2)
para("ผู้วิจัยนำข้อมูล 200 ค่าสังเกตที่ผ่านการรวบรวมและตรวจสอบแล้ว มาประมวลผลด้วยโปรแกรมสำเร็จรูปทางสถิติและ"
     "เศรษฐมิติ (เช่น STATA, EViews หรือ Python) กำหนดระดับนัยสำคัญทางสถิติที่ 0.05 โดยแบ่งการวิเคราะห์เป็น 5 ขั้นตอน")

heading("6.4.1 การวิเคราะห์สถิติเชิงพรรณนา (Descriptive Statistics)", H2, 3)
para("อธิบายลักษณะการกระจายตัวของตัวแปรแต่ละตัว ทั้งภาพรวมและจำแนกตามกลุ่ม ด้วยค่าเฉลี่ย (Mean) ส่วนเบี่ยงเบน"
     "มาตรฐาน (S.D.) ค่าสูงสุด (Maximum) และค่าต่ำสุด (Minimum) ตามแนวปฏิบัติมาตรฐานของการวิเคราะห์ข้อมูล"
     "ทางการเงิน (Deloof, 2003; Lind et al., 2012)")

heading("6.4.2 การทดสอบข้อตกลงเบื้องต้นของแบบจำลอง (Assumption Testing)", H2, 3)
para("ก่อนการประมาณค่าสมการถดถอย ผู้วิจัยทดสอบข้อตกลงเบื้องต้นเพื่อให้ผลการประมาณค่ามีความน่าเชื่อถือ ดังนี้")
runs_para([("1) ภาวะร่วมเส้นตรงพหุ (Multicollinearity): ", True, False),
           ("ตรวจด้วยค่า Variance Inflation Factor (VIF) ตามเกณฑ์ Hair et al. (2010) โดยค่า VIF ของทุกตัวแปรควรน้อยกว่า 10 "
            "ผู้วิจัยจะรายงานค่า VIF และเมทริกซ์สหสัมพันธ์ประกอบ โดยให้ความสำคัญกับความสัมพันธ์ระหว่าง DIV กับ SIZE เป็นพิเศษ "
            "เนื่องจากบริษัทที่กระจายความเสี่ยงต่างประเทศมักมีขนาดใหญ่", False, False)])
runs_para([("2) ความแปรปรวนของค่าคลาดเคลื่อนไม่คงที่ (Heteroskedasticity): ", True, False),
           ("ทดสอบด้วย Breusch–Pagan Test (Breusch & Pagan, 1979) หากค่า p-value มากกว่า 0.05 แสดงว่าความแปรปรวนคงที่ "
            "(Homoskedasticity) ทั้งนี้ หากพบปัญหาความแปรปรวนไม่คงที่ ผู้วิจัยจะแก้ไขด้วยการประมาณค่าความคลาดเคลื่อน"
            "มาตรฐานที่ทนทาน (Robust/Clustered Standard Errors)", False, False)])
runs_para([("3) สหสัมพันธ์ในตัวของข้อมูลแบบแผง (Autocorrelation): ", True, False),
           ("เนื่องจากข้อมูลเป็นแบบแผงที่มีมิติเวลา ผู้วิจัยใช้การทดสอบ Wooldridge Test สำหรับสหสัมพันธ์ในตัวของข้อมูลแบบแผง "
            "(Wooldridge, 2010) ซึ่งเหมาะสมกับโครงสร้าง Panel มากกว่าค่า Durbin–Watson ที่ออกแบบสำหรับอนุกรมเวลาชุดเดียว "
            "และแก้ไขผลกระทบด้วยการประมาณค่าความคลาดเคลื่อนมาตรฐานแบบคลัสเตอร์ตามบริษัท (Firm-clustered Standard Errors)", False, False)], space_after=8)

heading("6.4.3 การวิเคราะห์สหสัมพันธ์ (Correlation Analysis)", H2, 3)
para("ใช้สัมประสิทธิ์สหสัมพันธ์ของเพียร์สัน (Pearson Correlation Coefficient) เพื่อดูทิศทางและระดับความสัมพันธ์เชิงเส้น"
     "ระหว่างตัวแปรเป็นรายคู่ ทั้งตัวแปรตาม ตัวแปรอิสระ และตัวแปรควบคุม ค่าสัมประสิทธิ์ (r) มีค่าระหว่าง −1 ถึง 1 "
     "โดยใช้ประกอบการพิจารณาภาวะร่วมเส้นตรงพหุด้วย")

heading("6.4.4 การเลือกแบบจำลองข้อมูลแบบแผง (Panel Model Selection)", H2, 3)
para("การวิเคราะห์ข้อมูลแบบแผงมีแบบจำลอง 3 รูปแบบ ได้แก่ Pooled OLS, Fixed Effects (FE) และ Random Effects (RE) "
     "โดยทั่วไปใช้การทดสอบ Breusch–Pagan LM (เลือกระหว่าง Pooled OLS กับ RE) และ Hausman Test (Hausman, 1978; "
     "เลือกระหว่าง FE กับ RE)")
runs_para([("การตัดสินใจเชิงระเบียบวิธีของงานวิจัยนี้: ", True, True),
           ("เนื่องจากตัวแปรอิสระหลัก คือ DIV เป็นตัวแปรที่ไม่เปลี่ยนแปลงตามเวลา (Time-invariant) หากใช้แบบจำลอง "
            "Fixed Effects ตัวแปร DIV จะถูกกำจัดออกโดยอัตโนมัติ เพราะซ้อนทับกับผลกระทบเฉพาะบริษัท (Entity Effect) "
            "ทำให้ไม่สามารถประเมินและตีความสัมประสิทธิ์ของ DIV ซึ่งเป็นหัวใจของสมมติฐาน H1–H3 ได้ ", False, False),
           ("ด้วยเหตุนี้ ผู้วิจัยจึงเลือกใช้แบบจำลอง Random Effects (RE) เป็นแบบจำลองหลัก ", True, False),
           ("และรายงานผล Breusch–Pagan LM และ Hausman Test ประกอบเพื่อความครบถ้วน ทั้งนี้ หากผล Hausman ชี้ไปทาง FE "
            "ผู้วิจัยจะยังคงใช้ RE พร้อมระบุเหตุผลด้าน Time-invariance ของ DIV อย่างชัดเจน (ทางเลือกขั้นสูงอาจใช้ตัวประมาณค่า "
            "Hausman–Taylor ที่ประเมินตัวแปร Time-invariant ได้) การประมาณค่าทั้งหมดใช้ค่าความคลาดเคลื่อนมาตรฐานแบบ"
            "คลัสเตอร์ตามบริษัท", False, False)], space_after=8)

heading("6.4.5 การทดสอบสมมติฐานและการวิเคราะห์สมการถดถอย", H2, 3)
runs_para([("ส่วนที่ 1 — การเปรียบเทียบค่าเฉลี่ยระหว่างกลุ่ม: ", True, False),
           ("ใช้สถิติทดสอบค่าที (Independent-samples t-test) เปรียบเทียบ CCC, ROA และ ROE ระหว่างสองกลุ่ม โดยดำเนินการบน "
            "“ค่าเฉลี่ยระดับบริษัท” (n = 3 เทียบกับ n = 7) ไม่ใช่ระดับค่าสังเกตรายไตรมาส เพื่อหลีกเลี่ยงการละเมิดสมมติฐาน"
            "ความเป็นอิสระของข้อมูล (เนื่องจาก 20 ไตรมาสของบริษัทเดียวกันมีความสัมพันธ์กันเอง) ทั้งนี้การเปรียบเทียบที่คุม"
            "ตัวแปรอื่นแล้วจะพิจารณาจากสัมประสิทธิ์ DIV ในสมการถดถอยเป็นหลัก", False, False)])
runs_para([("ส่วนที่ 2 — การวิเคราะห์สมการถดถอยพหุคูณสำหรับข้อมูลแบบแผง: ", True, False),
           ("ตามกรอบแนวคิดของ Baltagi (2008) และ Wooldridge (2010) เพื่อทดสอบผลกระทบของการกระจายความเสี่ยงต่อวงจรเงินสด"
            "และความสามารถในการทำกำไร โดยควบคุมขนาดกิจการ โครงสร้างทางการเงิน ฤดูกาล และช่วงวิกฤต ดังสมการในหัวข้อ 6.5", False, False)], space_after=8)

# ---- 6.5 ----
heading("6.5 สมการที่ใช้ในการวิจัย", H1, 2)
para("แบบจำลองถดถอยเชิงพหุคูณสำหรับข้อมูลแบบแผง (Panel Data Regression) แยกตามตัวแปรตาม 3 สมการ ดังนี้")
para("สมการที่ 1 — ผลกระทบต่อวงจรเงินสด (CCC):", bold=True, space_after=3)
formula("CCCit = β0 + β1DIVit + β2SIZEit + β3LEVit + β4Q2t + β5Q3t + β6Q4t + β7COVIDit + εit")
para("สมการที่ 2 — ผลกระทบต่ออัตราผลตอบแทนต่อสินทรัพย์รวม (ROA):", bold=True, space_after=3)
formula("ROAit = β0 + β1DIVit + β2SIZEit + β3LEVit + β4Q2t + β5Q3t + β6Q4t + β7COVIDit + εit")
para("สมการที่ 3 — ผลกระทบต่ออัตราผลตอบแทนต่อส่วนของผู้ถือหุ้น (ROE):", bold=True, space_after=3)
formula("ROEit = β0 + β1DIVit + β2SIZEit + β3LEVit + β4Q2t + β5Q3t + β6Q4t + β7COVIDit + εit")
para("โดยกำหนดให้:", indent=False, bold=True, space_after=3)
for txt in [
    "i = ลำดับบริษัทตัวอย่าง (บริษัทที่ 1 ถึง 10)",
    "t = ระยะเวลารายไตรมาส (ไตรมาสที่ 1 ถึง 20)",
    "β0 = ค่าคงที่ (Intercept); β1–β7 = ค่าสัมประสิทธิ์การถดถอย",
    "DIV = ตัวแปรหุ่นการกระจายความเสี่ยงต่างประเทศ (1 = ต่างประเทศ, 0 = ในประเทศ)",
    "SIZE = ขนาดกิจการ (ln สินทรัพย์รวม); LEV = อัตราส่วนหนี้สินต่อทุน",
    "Q2, Q3, Q4 = ตัวแปรหุ่นฤดูกาล (ไตรมาส 1 เป็นฐาน); COVID = ตัวแปรหุ่นช่วงวิกฤตโรคระบาด",
    "εit = ค่าความคลาดเคลื่อน (ประมาณด้วยค่าความคลาดเคลื่อนมาตรฐานแบบคลัสเตอร์ตามบริษัท)",
]:
    bullet(txt)
runs_para([("การทดสอบความคงทน (Robustness Check): ", True, False),
           ("ผู้วิจัยจะประมาณค่าสมการทั้งชุดข้อมูลเต็ม (200 ค่าสังเกต) และชุดที่ตัดช่วงวิกฤตโรคระบาด เพื่อยืนยันว่าผลการศึกษา"
            "ไม่ได้ถูกครอบงำด้วยค่าผิดปกติในช่วงรายได้หดตัว และจะทดสอบด้วยตัวแปร DIV แบบต่อเนื่อง (สัดส่วนรายได้ต่างประเทศ) "
            "เป็นแบบจำลองทางเลือก", False, False)], space_after=8)

# ---- 6.6 ----
heading("6.6 ข้อจำกัดของงานวิจัย", H1, 2)
bullet("กลุ่มที่กระจายความเสี่ยงต่างประเทศมีเพียง 3 บริษัท (MINT, SHR, DUSIT) ทำให้สัมประสิทธิ์ DIV ถูกขับเคลื่อนด้วย"
       "ลักษณะเฉพาะของบริษัทเหล่านี้ และมีอำนาจการทดสอบทางสถิติ (Statistical Power) จำกัด")
bullet("ข้อมูลเป็นแผงขนาดเล็ก (N = 10, T = 20) คุณสมบัติเชิงเส้นกำกับ (Asymptotic) ของการทดสอบบางตัวจึงมีข้อจำกัด "
       "ผู้วิจัยจึงพึ่งพาค่าความคลาดเคลื่อนมาตรฐานแบบคลัสเตอร์เป็นหลัก")
bullet("ช่วงเวลาศึกษาคร่อมวิกฤตโรคระบาด (COVID-19) ซึ่งเป็นเหตุการณ์พิเศษที่กระทบธุรกิจโรงแรมรุนแรง แม้ควบคุมด้วย"
       "ตัวแปรหุ่นและการตัดช่วงวิกฤตแล้ว ผลการศึกษาบางส่วนยังอาจสะท้อนภาวะฟื้นตัวเฉพาะช่วง")
bullet("การจำแนกกลุ่มด้วยตัวแปรหุ่นอาจไม่สะท้อนระดับการกระจายความเสี่ยงที่แท้จริง ผู้วิจัยจึงเสริมด้วยสัดส่วนรายได้"
       "ต่างประเทศแบบต่อเนื่องเป็นการทดสอบความคงทน")

out = r"output/บทที่6_ระเบียบวิธีวิจัย_ฉบับแก้.docx"
doc.save(out)
print("saved:", out)
